#Import required libraries
import nltk
import os
from docx import Document
from bs4 import BeautifulSoup
from nltk import word_tokenize
from nltk import sent_tokenize
from nltk.sentiment.vader import SentimentIntensityAnalyzer

'''
This script will check the positive and negative meaning of the sentence and 
then separate the facts and consequences of negative sentence for Risk Factors and Business Sections of 10-K 
and 424B forms which were extracted using sectionalizing scripts.
'''

#Extracting risks and facts of the text
def get_risk_fact(text):
    lines = sent_tokenize(text)

    facts = []
    risks = []
    modals = set()

    for line in lines:
        st = set()
        text = word_tokenize(line)
        chapter = line
        sid = SentimentIntensityAnalyzer()
        ss = sid.polarity_scores(chapter)

        pos = ss["pos"]
        neg = ss["neg"]
        neu = ss["neu"]
        sent = "Positive: " + str(pos) + "  Negative: " + str(neg)

        for p_tags in nltk.pos_tag(text):
            st.add(p_tags[1])
			
			# Finding risks and facts 
            if "MD" == p_tags[1]:
                modals.add(p_tags[0])
        if "MD" in st:
            risks.append(line + "\n" + sent + "\n")
        else:
            facts.append(line + "\n" + sent + "\n")
    return facts, risks


path = "C:\\Users\\Vishal\\Desktop\\RiskFact\\Risk_Facts_separated"
os.chdir(path)
for root, dirs, files in os.walk(os.getcwd()):
    for file in files:
        file_ptr = open(file, "r", encoding="utf-8")
        file_data = file_ptr.read()
        file_ptr.close()
        soup = BeautifulSoup(file_data)
        # print(soup.text)
        text = soup.text
        text = text.replace("\n", " ")
        text = text.replace("  ", "")
		
		#Extracting risks and facts of the text
        facts, risks = get_risk_fact(text)

		# Storing risks and facts into a document for further analysis purpose
        risk_file = Document()
        risk_file.add_heading("Risks of " + file)
        for r in risks:
            risk_file.add_paragraph(r)
        risk_file.save(file[:-5] + '_RISK' + ".text")

        fact_file = Document()
        fact_file.add_heading("Facts of " + file)
        for f in facts:
            fact_file.add_paragraph(f)
        fact_file.save(file[:-5] + '_FACT' + ".text")

        # risk_file = open ("RISK" +file[:-5] + ".txt", 'w', encoding="utf-8" )
        #
        # risk_file.writelines(risks)
        # risk_file.close()
        #
        # fact_file = open("FACT" + file[:-5] + ".txt", 'w', encoding="utf-8")
        # fact_file.writelines(facts)
        # fact_file.close()